"""Ported OSWorld libreoffice_writer evaluator functions.

These are programmatic graders for our 21 OSWorld-Verified docx tasks.  Each
function returns a float in [0, 1] (pass = 1, fail = 0; some are fractional).

Source: https://github.com/xlang-ai/OSWorld/blob/main/desktop_env/evaluators/metrics/docs.py
Apache-2.0 license — ports adapted to our manifest's `evaluator.checks` schema.

Differences from upstream:
  - Removed verbose stdout prints / colored output.
  - Heavy deps (easyocr, skimage) are imported lazily; unavailable → return 0.
  - `find_default_font` is stubbed (operates on a LibreOffice config XML that
    doesn't exist in our headless code-execution sandbox).
  - `infeasible` semantic added: pass iff the working file is byte-identical
    to the source (the agent correctly refused to modify).
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import Any, Callable, Dict, List

logger = logging.getLogger("docx_metrics")

# All evaluators are wrapped to be tolerant: any unexpected exception → 0.0.


# ---------------------------------------------------------------------------
# compare_docx_files — content (paragraph) comparison with options
# ---------------------------------------------------------------------------

def compare_docx_files(file1: str, file2: str, **options) -> float:
    if not file1 or not file2:
        return 0.0

    ignore_blanks = options.get("ignore_blanks", True)
    ignore_case = options.get("ignore_case", False)
    ignore_order = options.get("ignore_order", False)
    content_only = options.get("content_only", False)
    fuzzy_match = options.get("fuzzy_match", False)
    delete_empty_lines = options.get("delete_empty_lines", False)

    try:
        from docx import Document
        doc1 = Document(file1)
        doc2 = Document(file2)
    except Exception as e:
        logger.debug(f"compare_docx_files load failed: {e}")
        return 0.0

    p1 = [p.text for p in doc1.paragraphs]
    p2 = [p.text for p in doc2.paragraphs]
    if ignore_order:
        p1, p2 = sorted(p1), sorted(p2)
    if delete_empty_lines:
        p1 = [p for p in p1 if p.strip()]
        p2 = [p for p in p2 if p.strip()]

    def _fuzzy_ratio(a: str, b: str) -> float:
        try:
            from rapidfuzz import fuzz
            return fuzz.ratio(a, b) / 100.0
        except Exception:
            # Fallback to a deterministic, dep-free Levenshtein-ish ratio
            from difflib import SequenceMatcher
            return SequenceMatcher(None, a, b).ratio()

    if content_only:
        text1 = re.sub(r"\s+", " ", "\n".join(p1)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(p2)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        return _fuzzy_ratio(text1, text2)

    if ignore_blanks:
        text1 = re.sub(r"\s+", " ", "\n".join(p1)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(p2)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        if fuzzy_match:
            return _fuzzy_ratio(text1, text2)
        return 1.0 if text1 == text2 else 0.0

    if len(p1) != len(p2):
        return 0.0
    if fuzzy_match:
        if not p1:
            return 1.0
        total = 0.0
        for a, b in zip(p1, p2):
            if ignore_case:
                a, b = a.lower(), b.lower()
            total += _fuzzy_ratio(a, b)
        return total / len(p1)
    for a, b in zip(p1, p2):
        if ignore_case:
            a, b = a.lower(), b.lower()
        if a != b:
            return 0.0
    return 1.0


# ---------------------------------------------------------------------------
# Table / image / line-spacing / font / structural checks
# ---------------------------------------------------------------------------

def compare_docx_tables(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    try:
        from docx import Document
        d1 = Document(file1)
        d2 = Document(file2)
    except Exception:
        return 0.0
    t1, t2 = d1.tables, d2.tables
    if len(t1) != len(t2):
        return 0.0
    for x, y in zip(t1, t2):
        if len(x.rows) != len(y.rows) or len(x.columns) != len(y.columns):
            return 0.0
        for i in range(len(x.rows)):
            for j in range(len(x.columns)):
                if x.cell(i, j).text.strip() != y.cell(i, j).text.strip():
                    return 0.0
    return 1.0


def compare_docx_images(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    try:
        from docx import Document
        from PIL import Image
        d1 = Document(file1)
        d2 = Document(file2)
    except Exception:
        return 0.0

    def _imgs(d):
        out = []
        for rel in d.part.rels.values():
            if "image" in rel.reltype:
                out.append(BytesIO(rel.target_part.blob))
        return out

    a, b = _imgs(d1), _imgs(d2)
    if len(a) != len(b):
        return 0.0
    for x, y in zip(a, b):
        try:
            if Image.open(x).tobytes() != Image.open(y).tobytes():
                return 0.0
        except Exception:
            return 0.0
    return 1.0


def compare_line_spacing(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    if compare_docx_files(file1, file2) == 0.0:
        return 0.0
    try:
        from docx import Document
        d1 = Document(file1)
        d2 = Document(file2)
    except Exception:
        return 0.0
    if len(d1.paragraphs) != len(d2.paragraphs):
        return 0.0
    for a, b in zip(d1.paragraphs, d2.paragraphs):
        if a.paragraph_format.line_spacing != b.paragraph_format.line_spacing:
            return 0.0
    return 1.0


def compare_subscript_contains(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    try:
        from docx import Document
        d1 = Document(file1)
        d2 = Document(file2)
    except Exception:
        return 0.0
    for a, b in zip(d1.paragraphs, d2.paragraphs):
        for ra, rb in zip(a.runs, b.runs):
            if ra.font.subscript and rb.font.subscript:
                return 1.0
    return 0.0


def check_tabstops(file1: str, file2: str, **kwargs) -> float:
    if not file1 or not file2:
        return 0.0
    try:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT
        d1 = Document(file1)
        d2 = Document(file2)
    except Exception:
        return 0.0

    p1 = [p for p in d1.paragraphs if p.text.strip()]
    p2 = [p for p in d2.paragraphs if p.text.strip()]
    if len(p1) != len(p2):
        return 0.0

    if kwargs.get("word_number_split_by_tabstop") is not None:
        number = kwargs["word_number_split_by_tabstop"]
        index = kwargs.get("index", 0)
        for p in p1:
            splits = p.text.split("\t")
            if len(splits) == 0 or index >= len(splits):
                return 0.0
            words = [w for w in re.split(r"\s", splits[index]) if w.strip()]
            if len(words) != number:
                return 0.0

    section = d2.sections[0]
    page_w = section.page_width - section.left_margin - section.right_margin

    def _ignore(t):
        return t.alignment == WD_TAB_ALIGNMENT.CLEAR or (
            t.alignment == WD_TAB_ALIGNMENT.LEFT and t.position == 0
        )

    minus = 0.0
    for a, b in zip(p1, p2):
        ta = [t for t in a.paragraph_format.tab_stops if not _ignore(t)]
        tb = [t for t in b.paragraph_format.tab_stops if not _ignore(t)]
        if len(ta) != len(tb):
            return 0.0
        diff = 0.0
        for ta_i, tb_i in zip(ta, tb):
            if ta_i.alignment != tb_i.alignment:
                return 0.0
            diff += abs(ta_i.position - tb_i.position)
        minus += diff / page_w
    return max(0.0, 1.0 - (minus / max(1, len(p1))))


# ---------------------------------------------------------------------------
# Single-file property assertions
# ---------------------------------------------------------------------------

def has_page_numbers_in_footers(file1: str, **_options) -> float:
    if not file1:
        return 0.0
    try:
        from docx import Document
        d = Document(file1)
    except Exception:
        return 0.0
    for section in d.sections:
        footer = section.footer
        if footer is None:
            return 0.0
        text = footer.paragraphs[0].text if footer.paragraphs else ""
        if not any(c.isdigit() for c in text):
            return 0.0
    return 1.0


def is_first_line_centered(file1: str, **_options) -> float:
    if not file1:
        return 0.0
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        d = Document(file1)
    except Exception:
        return 0.0
    if not d.paragraphs:
        return 0.0
    return 1.0 if d.paragraphs[0].paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER else 0.0


def contains_page_break(file1: str, **options) -> float:
    if not file1:
        return 0.0
    try:
        from docx import Document
        d = Document(file1)
    except Exception:
        return 0.0
    expected = options.get("page_break_count")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    count = 0
    for p in d.paragraphs:
        for run in p.runs:
            for br in run.element.findall(".//w:br", ns):
                t = br.attrib.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                )
                if t == "page":
                    count += 1
    if expected is not None and count != expected:
        return 0.0
    return 1.0 if count > 0 else 0.0


def compare_font_names(file1: str, **options) -> float:
    """Single-file evaluator. `options['font_name']` is the expected font."""
    if not file1:
        return 0.0
    try:
        from docx import Document
        d = Document(file1)
    except Exception:
        return 0.0
    expected = options.get("font_name")
    if expected is None:
        return 0.0
    for p in d.paragraphs:
        for run in p.runs:
            if run.font.name != expected:
                return 0.0
    return 1.0


def check_italic_font_size_14(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    if compare_docx_files(file1, file2) == 0.0:
        return 0.0
    try:
        from docx import Document
        d = Document(file1)
    except Exception:
        return 0.0
    for p in d.paragraphs:
        for run in p.runs:
            if run.italic:
                if run.font.size is None or run.font.size.pt != 14:
                    return 0.0
    return 1.0


def evaluate_strike_through_last_paragraph(file1: str, file2: str, **_options) -> float:
    if not file1 or not file2:
        return 0.0
    if compare_docx_files(file1, file2) == 0.0:
        return 0.0
    try:
        from docx import Document
        d = Document(file1)
    except Exception:
        return 0.0
    if not d.paragraphs:
        return 0.0
    last = d.paragraphs[-1]
    if not last.runs:
        return 0.0
    for run in last.runs:
        if not run.font.strike:
            return 0.0
    return 1.0


def evaluate_colored_words_in_tables(file1: str, file2: str, **options) -> float:
    """Vowel-initial words red, others blue (CIE delta-E 2000 distance check)."""
    if not file1 or not file2:
        return 0.0
    if compare_docx_files(file1, file2) == 0.0:
        return 0.0
    try:
        from docx import Document
        from docx.shared import RGBColor
    except Exception:
        return 0.0
    try:
        # skimage is optional; if missing, fall back to a strict RGB equality check.
        from skimage.color import deltaE_ciede2000, rgb2lab  # type: ignore
        _have_skimage = True
    except Exception:
        _have_skimage = False

    threshold = options.get("threshold", 3.5)

    def _color_diff(a, b):
        if _have_skimage:
            la = rgb2lab([a[0] / 255.0, a[1] / 255.0, a[2] / 255.0])
            lb = rgb2lab([b[0] / 255.0, b[1] / 255.0, b[2] / 255.0])
            return float(deltaE_ciede2000(la, lb))
        return 0.0 if (a == b) else 999.0  # fallback: equality only

    try:
        d = Document(file1)
    except Exception:
        return 0.0

    red = RGBColor(255, 0, 0)
    blue = RGBColor(0, 0, 255)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        word = run.text
                        if not word:
                            continue
                        rgb = run.font.color.rgb
                        if rgb is None:
                            return 0.0
                        first = word[0].lower()
                        if first in "aeiou":
                            if _color_diff((rgb[0], rgb[1], rgb[2]), (255, 0, 0)) > threshold:
                                return 0.0
                        else:
                            if _color_diff((rgb[0], rgb[1], rgb[2]), (0, 0, 255)) > threshold:
                                return 0.0
    return 1.0


# ---------------------------------------------------------------------------
# Multi-file evaluator — compare_unique_train_records
# ---------------------------------------------------------------------------

def compare_unique_train_records(processed_file: str, expected_files: List[str], **_options) -> float:
    """expected_files = [gold, initial].  Returns 1 iff:
      - Every processed line appeared in initial
      - No duplicate train_ids in processed
      - Set of train_ids matches gold
      - Line counts match
    """
    if not processed_file or not isinstance(expected_files, list) or len(expected_files) < 2:
        return 0.0
    try:
        from docx import Document
    except Exception:
        return 0.0

    gold, initial = expected_files[0], expected_files[1]
    if not gold or not initial:
        return 0.0

    def _lines_and_ids(path):
        try:
            d = Document(path)
            lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            ids = [ln.split(",")[1].strip() for ln in lines if len(ln.split(",")) == 4]
            return lines, ids
        except Exception:
            return None, None

    proc_l, proc_ids = _lines_and_ids(processed_file)
    if proc_l is None:
        return 0.0
    gold_l, gold_ids = _lines_and_ids(gold)
    if gold_l is None:
        return 0.0
    init_l, _ = _lines_and_ids(initial)
    if init_l is None:
        return 0.0

    if not set(proc_l).issubset(set(init_l)):
        return 0.0
    if len(proc_ids) != len(set(proc_ids)):
        return 0.0
    if set(proc_ids) != set(gold_ids):
        return 0.0
    if len(proc_l) != len(gold_l):
        return 0.0
    return 1.0


# ---------------------------------------------------------------------------
# Special: infeasible / unsupported
# ---------------------------------------------------------------------------

def infeasible(file1: str, source_file: str | None = None, **_options) -> float:
    """Pass iff the agent's file is byte-identical to the source — i.e. the
    agent correctly recognized the task as impossible and didn't modify."""
    if not file1 or not source_file:
        return 0.0
    try:
        with open(file1, "rb") as f:
            a = f.read()
        with open(source_file, "rb") as f:
            b = f.read()
        return 1.0 if a == b else 0.0
    except Exception:
        return 0.0


def find_default_font(file1: str, **_options) -> float:
    """Stub — operates on LibreOffice's XML config which doesn't exist in our
    headless code-execution sandbox.  Always returns 0 so the docx grader can
    fall back to the diff layer."""
    return 0.0


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

# Two argument styles:
#   binary: f(working_file, gold_file, **options)
#   single: f(working_file, **options)            -- options carry the rule
#   multi:  f(working_file, expected_files: list, **options)

_BINARY: Dict[str, Callable] = {
    "compare_docx_files": compare_docx_files,
    "compare_docx_tables": compare_docx_tables,
    "compare_docx_images": compare_docx_images,
    "compare_line_spacing": compare_line_spacing,
    "compare_subscript_contains": compare_subscript_contains,
    "check_tabstops": check_tabstops,
    "check_italic_font_size_14": check_italic_font_size_14,
    "evaluate_strike_through_last_paragraph": evaluate_strike_through_last_paragraph,
    "evaluate_colored_words_in_tables": evaluate_colored_words_in_tables,
}

_SINGLE: Dict[str, Callable] = {
    "has_page_numbers_in_footers": has_page_numbers_in_footers,
    "is_first_line_centered": is_first_line_centered,
    "contains_page_break": contains_page_break,
    "compare_font_names": compare_font_names,
    "find_default_font": find_default_font,
}

_MULTI: Dict[str, Callable] = {
    "compare_unique_train_records": compare_unique_train_records,
}


def run_check(
    func: str,
    working_file: str,
    expected_files: List[str],
    options: Dict[str, Any],
    source_file: str | None = None,
) -> float:
    """Run one OSWorld evaluator check.  Returns float in [0, 1]."""
    try:
        opts = options or {}
        if func == "infeasible":
            return infeasible(working_file, source_file=source_file)
        if func in _BINARY:
            gold = expected_files[0] if expected_files else ""
            return float(_BINARY[func](working_file, gold, **opts))
        if func in _SINGLE:
            return float(_SINGLE[func](working_file, **opts))
        if func in _MULTI:
            return float(_MULTI[func](working_file, expected_files, **opts))
    except Exception as e:
        logger.debug(f"run_check {func} crashed: {e}")
        return 0.0
    return 0.0


def run_evaluator(
    conj: str,
    checks: List[Dict[str, Any]],
    working_file: str,
    source_file: str | None = None,
) -> float:
    """Combine `conj` ('and' | 'or') over a list of checks.

    Each check dict has: func (str), options (dict), expected_files (list[str]).
    `and`: min over scores.  `or`: max over scores.
    """
    if not checks:
        return 0.0
    scores = [
        run_check(
            func=c.get("func", ""),
            working_file=working_file,
            expected_files=c.get("expected_files") or [],
            options=c.get("options") or {},
            source_file=source_file,
        )
        for c in checks
    ]
    if conj == "or":
        return max(scores)
    return min(scores)
